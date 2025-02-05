import os
from flask import Flask, request, render_template, send_file, url_for, send_from_directory
from werkzeug.utils import secure_filename
import fitz 
import shutil
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import pdfplumber
import pandas as pd
from openpyxl import Workbook
from moviepy import AudioFileClip
from docx import Document
import fitz
from fpdf import FPDF
from PIL import Image
import PyPDF2

app = Flask(__name__)

UPLOAD_FOLDER = os.path.join(os.getcwd(), "uploads")
OUTPUT_FOLDER = os.path.join(os.getcwd(), "converted")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 1024 * 1024 * 1024
app.config["TEMPLATES_AUTO_RELOAD"] = True

def is_valid_extension(filename, valid_extensions):
    return filename.lower().endswith(tuple(valid_extensions))

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/politica-privacidade')
def politica_privacidade():
    return render_template('politica_privacidade.html')

@app.route('/termos')
def termos_de_uso():
    return render_template('termos.html')


@app.route('/conversao_pdfword')
def conversao_pdfword():
    return render_template('indexpdfword.html')

@app.route('/conversao_mp4mp3')
def conversao_mp4mp3():
    return render_template('indexmp4mp3.html')

@app.route('/conversao_excelorpdf')
def conversao_excelorpdf():
    return render_template('indexexcelpdf.html')

@app.route('/conversao_txtpdf')
def conversao_txtpdf():
    return render_template('indexpdftxt.html')

@app.route('/conversao_pdfimages')
def conversao_pdfimages():
    return render_template('indexpdfimagens.html')

@app.route('/conversao_pngparaico')
def conversao_pngparaico():
    return render_template('indexpngtoico.html')
 
@app.route("/convert_image", methods=["POST"])
def convert_image():
    file = request.files.get("file")

    if not file:
        return "Nenhum arquivo enviado.", 400

    input_filename = secure_filename(file.filename)
    if not is_valid_extension(input_filename, [".pdf"]):
        return "Formato de arquivo inválido. Apenas PDF é permitido.", 400

    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

    output_folder = os.path.join(app.config["OUTPUT_FOLDER"], os.path.splitext(input_filename)[0])
    os.makedirs(output_folder, exist_ok=True)

    output_format = request.form.get("outputFormat", "PNG").upper()

    if output_format not in ["PNG", "JPG"]:
        return "Formato de saída inválido. Apenas PNG ou JPG são permitidos.", 400

    try:
        convert_pdf_to_images(input_path, output_folder, output_format)
    except Exception as e:
        return str(e), 500

    download_link = url_for("download_image", folder=os.path.basename(output_folder))
    return render_template("indexpdfimagens.html", download_link=download_link)

@app.route("/convert_pngparaico", methods=["POST"])
def convert_png_to_ico_route():
    """Converte arquivos PNG para ICO."""
    file = request.files.get("file")
    icon_size = request.form.get("iconSize", "256")

    if not file:
        return "Nenhum arquivo enviado.", 400

    input_filename = secure_filename(file.filename)

    
    valid_extensions = [".png"]
    if not is_valid_extension(input_filename, valid_extensions):
        return f"Erro: O arquivo deve ser {', '.join(valid_extensions).upper()}.", 400

   
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

  
    output_folder = app.config["OUTPUT_FOLDER"]
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    
    output_filename = f"converted_{os.path.splitext(input_filename)[0]}.ico"
    output_path = os.path.join(output_folder, output_filename)

    try:
       
        convert_png_to_ico(input_path, output_path, int(icon_size))
    except Exception as e:
        return f"Erro inesperado: {str(e)}", 500

    
    zip_filename = f"{os.path.splitext(output_filename)[0]}.zip"
    zip_path = os.path.join(output_folder, zip_filename)

    try:
       
        shutil.make_archive(zip_path.replace('.zip', ''), 'zip', output_folder, output_filename)
        download_link = url_for("download_image", filename=zip_filename)
        return render_template("indexpngtoico.html", download_link=download_link)
    except Exception as e:
        return f"Erro ao criar o arquivo ZIP: {str(e)}", 500


@app.route("/download_image/<filename>")
def download_image(filename):
    """Baixa o arquivo ZIP com o arquivo ICO convertido."""
   
    output_folder = app.config["OUTPUT_FOLDER"]
    file_path = os.path.join(output_folder, filename)

   
    if not os.path.exists(file_path):
        return "Arquivo não encontrado.", 404

    
    return send_file(file_path, as_attachment=True)


def convert_png_to_ico(input_path, output_path, icon_size):
    """Converte um arquivo PNG para ICO com o tamanho especificado."""
    try:
        with Image.open(input_path) as img:
            img = img.resize((icon_size, icon_size), Image.Resampling.LANCZOS)
            img.save(output_path, format="ICO")
    except Exception as e:
        raise Exception(f"Erro ao converter PNG para ICO: {str(e)}")


def convert_pdf_to_images(input_path, output_folder, output_format):
    try:
        pdf_document = fitz.open(input_path)
        for page_number in range(len(pdf_document)):
            page = pdf_document.load_page(page_number)
            pix = page.get_pixmap()

            output_file = os.path.join(output_folder, f"page_{page_number + 1}.{output_format.lower()}")
            pix.save(output_file)
        pdf_document.close()
    except Exception as e:
        raise Exception(f"Erro ao converter PDF para imagens: {str(e)}")
    
@app.route("/convert_file", methods=["POST"])
def convert_file():
    file = request.files.get("file")
    conversion_type = request.form.get("conversionType")

    if not file:
        return "Nenhum arquivo enviado.", 400

    input_filename = secure_filename(file.filename)
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

    output_filename = None
    output_path = None

    try:
        if conversion_type == "pdf-to-excel" and is_valid_extension(input_filename, [".pdf"]):
            output_filename = f"converted_{os.path.splitext(input_filename)[0]}.xlsx"
            output_path = os.path.join(app.config["OUTPUT_FOLDER"], output_filename)
            convert_pdf_to_excel(input_path, output_path)

        elif conversion_type == "excel-to-pdf" and is_valid_extension(input_filename, [".xlsx"]):
            output_filename = f"converted_{os.path.splitext(input_filename)[0]}.pdf"
            output_path = os.path.join(app.config["OUTPUT_FOLDER"], output_filename)
            convert_excel_to_pdf(input_path, output_path)

        else:
            return "Erro: Tipo de conversão inválido ou formato de arquivo errado.", 400

    except Exception as e:
        return f"Erro ao converter arquivo: {str(e)}", 500

    download_link = url_for("download_file", filename=os.path.basename(output_path))
    return render_template("indexexcelpdf.html", download_link=download_link)

@app.route("/convert_video_to_audio", methods=["POST"])
def convert_video_to_audio_route():
    file = request.files.get("file")

    if not file:
        return "Nenhum arquivo enviado.", 400

    input_filename = secure_filename(file.filename)
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

    output_filename = f"converted_{os.path.splitext(input_filename)[0]}.mp3"
    output_path = os.path.join(app.config["OUTPUT_FOLDER"], output_filename)

    try:
        if not is_valid_extension(input_filename, [".mp4", ".m4v"]):
            return "Erro: O arquivo deve ser MP4 ou M4V.", 400
        convert_video_to_audio(input_path, output_path)
    except Exception as e:
        return f"Erro ao converter vídeo para áudio: {str(e)}", 500

    download_link = url_for("download_file", filename=os.path.basename(output_path))
    return render_template("indexmp4mp3.html", download_link=download_link)

@app.route("/download/<filename>")
def download_file(filename):
    file_path = os.path.join(app.config["OUTPUT_FOLDER"], filename)
    if not os.path.exists(file_path):
        return "Arquivo não encontrado.", 404
    return send_file(file_path, as_attachment=True)

@app.route("/convert_pdf_to_excel", methods=["POST"])
def convert_pdf_to_excel_route():
    file = request.files.get("file")

    if not file:
        return "Nenhum arquivo enviado.", 400

    input_filename = secure_filename(file.filename)
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

    output_filename = f"converted_{os.path.splitext(input_filename)[0]}.xlsx"
    output_path = os.path.join(app.config["OUTPUT_FOLDER"], output_filename)

    try:
        if not is_valid_extension(input_filename, [".pdf"]):
            return "Erro: O arquivo deve ser um PDF.", 400
        convert_pdf_to_excel(input_path, output_path)
    except Exception as e:
        return f"Erro ao converter PDF para Excel: {str(e)}", 500

    download_link = url_for("download_file", filename=os.path.basename(output_path))
    return render_template("indexexcelpdf.html", download_link=download_link)

@app.route("/convert_image", methods=["POST"])
def convert_images():
    file = request.files.get("file")

    if not file:
        return "Nenhum arquivo enviado.", 400

    input_filename = secure_filename(file.filename)
    if not is_valid_extension(input_filename, [".pdf"]):
        return "Formato de arquivo inválido. Apenas PDF é permitido.", 400

    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

    output_folder = os.path.join(app.config["OUTPUT_FOLDER"], os.path.splitext(input_filename)[0])
    os.makedirs(output_folder, exist_ok=True)

    output_format = request.form.get("outputFormat", "PNG").upper()

    if output_format not in ["PNG", "JPG"]:
        return "Formato de saída inválido. Apenas PNG ou JPG são permitidos.", 400

    try:
        convert_pdf_to_images(input_path, output_folder, output_format)
    except Exception as e:
        return str(e), 500

    download_link = url_for("download_folder", folder=os.path.basename(output_folder))
    return render_template("indexpdfimagens.html", download_link=download_link)

@app.route("/download/<folder>")
def download_folder(folder):
    folder_path = os.path.join(app.config["OUTPUT_FOLDER"], folder)
    if not os.path.exists(folder_path):
        return "Pasta não encontrada.", 404

    zip_path = f"{folder_path}.zip"
    if not os.path.exists(zip_path):
        shutil.make_archive(folder_path, 'zip', folder_path)

    return send_file(zip_path, as_attachment=True)

def convert_pdf_to_images(input_path, output_folder, output_format):
    try:
        pdf_document = fitz.open(input_path)
        for page_number in range(len(pdf_document)):
            page = pdf_document.load_page(page_number)
            pix = page.get_pixmap()

            output_file = os.path.join(output_folder, f"page_{page_number + 1}.{output_format.lower()}")
            pix.save(output_file)
        pdf_document.close()
    except Exception as e:
        raise Exception(f"Erro ao converter PDF para imagens: {str(e)}")

@app.route("/convert_excel_to_pdf", methods=["POST"])
def convert_excel_to_pdf_route():
    file = request.files.get("file")

    if not file:
        return "Nenhum arquivo enviado.", 400

    input_filename = secure_filename(file.filename)
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

    output_filename = f"converted_{os.path.splitext(input_filename)[0]}.pdf"
    output_path = os.path.join(app.config["OUTPUT_FOLDER"], output_filename)

    try:
        if not is_valid_extension(input_filename, [".xlsx"]):
            return "Erro: O arquivo deve ser um Excel.", 400
        convert_excel_to_pdf(input_path, output_path)
    except Exception as e:
        return f"Erro ao converter Excel para PDF: {str(e)}", 500

    download_link = url_for("download_file", filename=os.path.basename(output_path))
    return render_template("indexexcelpdf.html", download_link=download_link)

@app.route("/converterpdf-txt", methods=["POST"])
def converter_pdf_txt():
    """Converte arquivos entre PDF e TXT com base no tipo selecionado."""
    file = request.files.get("file")
    conversion_type = request.form.get("conversionType")

    if not file or not conversion_type:
        return "Arquivo ou tipo de conversão não selecionado.", 400

    input_filename = secure_filename(file.filename)
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

    output_filename = f"converted_{os.path.splitext(input_filename)[0]}"
    output_path = None

    try:
        if conversion_type == "pdf-to-txt":
            if not is_valid_extension(input_filename, [".pdf"]):
                return "Erro: Para PDF para TXT, o arquivo enviado deve ser um PDF.", 400
            output_path = os.path.join(app.config["UPLOAD_FOLDER"], output_filename + ".txt")
            convert_pdf_to_txt(input_path, output_path)
        elif conversion_type == "txt-to-pdf":
            if not is_valid_extension(input_filename, [".txt"]):
                return "Erro: Para TXT para PDF, o arquivo enviado deve ser um TXT.", 400
            output_path = os.path.join(app.config["UPLOAD_FOLDER"], output_filename + ".pdf")
            convert_txt_to_pdf(input_path, output_path)
        else:
            return "Tipo de conversão inválido.", 400
    except Exception as e:
        return f"Erro inesperado: {str(e)}", 500

    return send_file(output_path, as_attachment=True, download_name=os.path.basename(output_path), mimetype="application/octet-stream")

@app.route("/download/<filename>")
def download_filetxt(filename):
    """Serve o arquivo convertido para download."""
    try:
        directory = app.config["UPLOAD_FOLDER"]
        return send_file(os.path.join(directory, filename), as_attachment=True)
    except Exception as e:
        return f"Erro ao tentar fazer o download: {str(e)}", 500

def convert_pdf_to_txt(input_path, output_path):
    """Converte PDF para TXT."""
    try:
        with open(output_path, "w", encoding="utf-8") as txt_file:
            with open(input_path, "rb") as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    txt_file.write(page.extract_text() + "\n")
    except Exception as e:
        raise Exception(f"Erro ao converter PDF para TXT: {str(e)}")

def convert_txt_to_pdf(input_path, output_path):
    """Converte TXT para PDF."""
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        with open(input_path, "r", encoding="utf-8") as txt_file:
            for line in txt_file:
                pdf.multi_cell(0, 10, line)
        pdf.output(output_path)
    except Exception as e:
        raise Exception(f"Erro ao converter TXT para PDF: {str(e)}")


@app.route("/converterpdfword", methods=["POST"])
def converter_pdf_e_word():
    """Converte arquivos entre PDF e Word com base no tipo selecionado."""
    file = request.files.get("file")
    conversion_type = request.form.get("conversionType")

    if not file or not conversion_type:
        return "Arquivo ou tipo de conversão não selecionado.", 400

    input_filename = secure_filename(file.filename)
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

    output_filename = f"converted_{os.path.splitext(input_filename)[0]}"
    output_path = None

    try:
        if conversion_type == "pdf-to-word":
            if not is_valid_extension(input_filename, [".pdf"]):
                return "Erro: Para PDF para Word, o arquivo enviado deve ser um PDF.", 400
            output_path = os.path.join(app.config["UPLOAD_FOLDER"], output_filename + ".docx")
            convert_pdf_to_word(input_path, output_path)
        elif conversion_type == "word-to-pdf":
            if not is_valid_extension(input_filename, [".docx"]):
                return "Erro: Para Word para PDF, o arquivo enviado deve ser um DOCX.", 400
            output_path = os.path.join(app.config["UPLOAD_FOLDER"], output_filename + ".pdf")
            convert_word_to_pdf(input_path, output_path)
        else:
            return "Tipo de conversão inválido.", 400
    except Exception as e:
        return f"Erro inesperado: {str(e)}", 500

    return send_file(output_path, as_attachment=True, download_name=os.path.basename(output_path))

def convert_pdf_to_word(input_path, output_path):
    try:
        doc = Document()
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
        doc.save(output_path)
    except Exception as e:
        raise Exception(f"Erro ao converter PDF para Word: {str(e)}")

def convert_word_to_pdf(input_path, output_path):
    try:
        doc = Document(input_path)
        c = canvas.Canvas(output_path, pagesize=letter)

        width, height = letter
        y = height - 50
        line_height = 12

        for paragraph in doc.paragraphs:
            for line in paragraph.text.split("\n"):
                if y < 50:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, line)
                y -= line_height

        c.save()
    except Exception as e:
        raise Exception(f"Erro ao converter Word para PDF: {str(e)}")

def convert_video_to_audio(input_path, output_path):
    try:
        audio_clip = AudioFileClip(input_path)
        audio_clip.write_audiofile(output_path)
        audio_clip.close()
    except Exception as e:
        raise Exception(f"Erro ao converter vídeo para áudio: {str(e)}")

def convert_pdf_to_excel(input_path, output_path):
    try:
        all_data = []
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        all_data.extend(table)
                else:
                    lines = page.extract_text().split("\n")
                    all_data.append(lines)

        df = pd.DataFrame(all_data)
        df.dropna(how="all", inplace=True)
        df.dropna(axis=1, how="all", inplace=True)
        df.to_excel(output_path, index=False, header=False)
    except Exception as e:
        raise Exception(f"Erro ao converter PDF para Excel: {str(e)}")

def convert_excel_to_pdf(input_path, output_path):
    try:
        wb = pd.ExcelFile(input_path)
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        y = height - 50
        line_height = 12

        for sheet in wb.sheet_names:
            df = wb.parse(sheet)
            df = df.dropna(how="all", axis=0)
            df = df.dropna(how="all", axis=1)

            for _, row in df.iterrows():
                row_text = ", ".join([str(value) for value in row.values if pd.notna(value)])
                if y < 50:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, row_text)
                y -= line_height

        c.save()
    except Exception as e:
        raise Exception(f"Erro ao converter Excel para PDF: {str(e)}")

@app.route("/", methods=["GET"])
def index():
    return render_template("index.html", download_link=None)

def is_valid_extension(filename, valid_extensions):
    return any(filename.lower().endswith(ext) for ext in valid_extensions)

@app.route("/convert_image", methods=["POST"])
def convert():
    file = request.files.get("file")

    if not file:
        return "Erro: Arquivo não selecionado.", 400

    input_filename = secure_filename(file.filename)
    if not is_valid_extension(input_filename, [".pdf"]):
        return "Erro: O arquivo enviado deve ser um PDF.", 400

    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

    if not os.path.exists(input_path):
        return "Erro: Arquivo não salvo corretamente.", 400

    output_folder = os.path.join(app.config["UPLOAD_FOLDER"], os.path.splitext(input_filename)[0])
    os.makedirs(output_folder, exist_ok=True)

    output_format = request.form.get("outputFormat", "PNG").upper()

    if output_format not in ["PNG", "JPG"]:
        return "Erro: Formato inválido. Apenas PNG e JPG são suportados.", 400

    try:
        convert_pdf_to_images(input_path, output_folder, output_format)
    except Exception:
        return "Erro ao converter o PDF.", 500

    output_folder_name = os.path.basename(output_folder)
    download_link = url_for("download_folder", folder=output_folder_name)
    return render_template("indexpdfimagens.html", download_link=download_link)

def convert_pdf_to_images(input_path, output_folder, output_format):
    try:
        pdf_document = fitz.open(input_path)
        for page_number in range(len(pdf_document)):
            page = pdf_document.load_page(page_number)
            pix = page.get_pixmap()

            output_file = os.path.join(output_folder, f"page_{page_number + 1}.{output_format.lower()}")
            pix.save(output_file)
        pdf_document.close()
    except Exception as e:
        raise Exception(f"Erro ao converter PDF para imagens: {str(e)}")
    
if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=25566)
