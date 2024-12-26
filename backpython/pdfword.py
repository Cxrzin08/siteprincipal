import os
from flask import Flask, request, render_template, send_file, url_for
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import pdfplumber

app = Flask(__name__)

UPLOAD_FOLDER = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config = {
    "UPLOAD_FOLDER": UPLOAD_FOLDER,
    "MAX_CONTENT_LENGTH": 10 * 1024 * 1024 
}

def is_valid_extension(filename, valid_extensions):
    """Verifica se o arquivo possui uma extensão válida."""
    return any(filename.lower().endswith(ext) for ext in valid_extensions)

@app.route("/", methods=["GET"])
def index():
    """Página inicial com formulário de upload."""
    return render_template("indexpdfword.html", download_link=None)

@app.route("/convert", methods=["POST"])
def convert():
    """Converte arquivos entre PDF e Word com base no tipo selecionado."""
    file = request.files.get("file")
    conversion_type = request.form.get("conversionType")

    if not file or not conversion_type:
        return "Arquivo ou tipo de conversão não selecionado.", 400

    input_filename = secure_filename(file.filename)
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], input_filename)
    file.save(input_path)

    output_filename = f"converted_{os.path.splitext(input_filename)[0]}"
    output_path = None

    try:
        if conversion_type == "pdf-to-word":
            if not is_valid_extension(input_filename, [".pdf"]):
                return "Erro: Para PDF para Word, o arquivo enviado deve ser um PDF.", 400
            output_path = os.path.join(app.config["UPLOAD_FOLDER"], output_filename + ".docx")
            convert_pdf_to_word(input_path, output_path)
        elif conversion_type == "word-to-pdf":
            if not is_valid_extension(input_filename, [".docx"]):
                return "Erro: Para Word para PDF, o arquivo enviado deve ser um DOCX.", 400
            output_path = os.path.join(app.config["UPLOAD_FOLDER"], output_filename + ".pdf")
            convert_word_to_pdf(input_path, output_path)
        else:
            return "Tipo de conversão inválido.", 400
    except Exception as e:
        return f"Erro inesperado: {str(e)}", 500

    download_link = url_for("download_file", filename=os.path.basename(output_path))
    return render_template("indexpdfword.html", download_link=download_link)

@app.route("/download/<filename>")
def download_file(filename):
    """Baixa o arquivo convertido."""
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    if not os.path.exists(file_path):
        return "Arquivo não encontrado.", 404
    return send_file(file_path, as_attachment=True)

def convert_pdf_to_word(input_path, output_path):
    """Converte PDF para Word com melhor extração de texto usando pdfplumber."""
    try:
        doc = Document()
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
        doc.save(output_path)
    except Exception as e:
        raise Exception(f"Erro ao converter PDF para Word: {str(e)}")

def convert_word_to_pdf(input_path, output_path):
    """Converte Word para PDF com formatação melhorada."""
    try:
        doc = Document(input_path)
        c = canvas.Canvas(output_path, pagesize=letter)

        width, height = letter
        y = height - 50
        line_height = 12

        for paragraph in doc.paragraphs:
            for line in paragraph.text.split("\n"):
                if y < 50:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, line)
                y -= line_height

        c.save()
    except Exception as e:
        raise Exception(f"Erro ao converter Word para PDF: {str(e)}")

if __name__ == "__main__":
    app.run(debug=True)
